import sys
import traceback
import requests
from datetime import datetime, timedelta, timezone
import re
import os
import boto3
from botocore.exceptions import ClientError
import json
from docx import Document
from presidio_analyzer import AnalyzerEngine
from presidio_analyzer.nlp_engine import SpacyNlpEngine
import spacy
import asyncio
import psycopg2
from urllib.parse import urlparse
from tqdm import tqdm as tqdm
import logging
import io

secret_name = os.environ['secret_name']
region_name = os.environ['region_name']
# ======================== LOGGER ============================
log_capture_string = io.StringIO()
ch = logging.StreamHandler(log_capture_string)
ch.setLevel(logging.DEBUG)

formatter = logging.Formatter('%(asctime)s %(levelname)-8s %(message)s', datefmt='%Y-%m-%d %H:%M:%S')
ch.setFormatter(formatter)

# Get logger
logger = logging.getLogger()
logger.setLevel(logging.INFO)
logger.addHandler(ch)

# =====================================================
# RETRIEVE SECRETS
# =====================================================
def get_secret(secret_name, region_name):
    session = boto3.session.Session()
    client = session.client("secretsmanager", region_name=region_name)
    secret = json.loads(client.get_secret_value(SecretId=secret_name)["SecretString"])
    return json.loads(secret['configPythonSecrets'])

config = get_secret(secret_name, region_name)

region_name = config['awsRegion']
db_conf = config['database']
llm_model = config['bedrock']['llm']
embedding_model = config['bedrock']['embedding']

db_config = {
    "dbname": db_conf['name'],
    "user": db_conf['user'],
    "password": db_conf['password'],
    "host": db_conf['host'],
    "port": db_conf['port'],
}

TENANT_ID = config['sharePoint']['tenantId']
CLIENT_ID = config['sharePoint']['clientId']
CLIENT_SECRET = config['sharePoint']['clientSecret']
site_url = config['config']['sharePoint']['teamsTranscript']
HOST_NAME = config['sharePoint']['baseUrl']
 
# Time filter: 3 years ago
three_years_ago = datetime.now(timezone.utc) - timedelta(days=3*365)
# Time filter: 24 hrs ago
last_24_hours = datetime.now(timezone.utc) - timedelta(hours=24)
 
#================ UTILS =================
 
def _sql_(query):
    # db_config = {  
    #     'dbname': db_name,  
    #     'user': db_user,  
    #     'password': db_password,  
    #     'host': db_host,  
    #     'port': db_port
    # }
    conn = psycopg2.connect(**db_config)
    cursor = conn.cursor()
 
    cursor.execute(query)
    out = cursor.fetchall()
    #print(f"SQL Query executed: {query} , Got {out}")
    if conn:
            cursor.close()
            conn.close()
    return out
 
 
# region_name = "eu-west-1"
# secret_name = "manual_lambda_secret"
 
# configPythonSecrets = get_secret(secret_name, region_name)
 
#=========================================================================================
#=========================================================================================
 
async def bedrock_invoke_model_async(prompt , region_name = region_name, llm_model = llm_model, max_tokens=800, temperature=0.3):
    """Invoke Bedrock model asynchronously."""
    client = boto3.client("bedrock-runtime", region_name=region_name)
    body = {
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": max_tokens,
        "temperature": temperature,
        "messages": [{"role": "user", "content": [{"type": "text", "text": prompt}]}],
    }
    def _invoke():
        response = client.invoke_model(modelId=llm_model, body=json.dumps(body))
        model_response = json.loads(response["body"].read())
        return model_response["content"][0]["text"].strip()
    return await asyncio.to_thread(_invoke)
 
def _get_embeddings(input_text, embedding_model = embedding_model, region_name = region_name):
 
    client = boto3.client("bedrock-runtime", region_name=region_name)
    native_request = {"inputText": input_text}
    request = json.dumps(native_request)
    response = client.invoke_model(modelId=embedding_model, body=request)
    model_response = json.loads(response["body"].read())
    embedding = model_response["embedding"]
    return embedding
 
#================ PII REMOVAL ==================
 
EMAIL_PATTERN = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", re.IGNORECASE)
PHONE_PATTERN = re.compile(r"\b\d{10}\b|\b\d{5}[ -]?\d{5}\b")
CREDIT_CARD_PATTERN = re.compile(r"\b(?:\d{4}[ -]?){3}\d{4}\b")  # handles 1234 5678 9012 3456 pattern
BANK_ACC_PATTERN = re.compile(r"\b\d{11,20}\b")
 
class LoadedSpacyNlpEngine(SpacyNlpEngine):
    def __init__(self, loaded_spacy_model):
        super().__init__()
        self.nlp = {"en": loaded_spacy_model}
 
# Load a model a-priori
nlp = spacy.load("./en_core_web_md-3.8.0")
 
# Pass the loaded model to the new LoadedSpacyNlpEngine
loaded_nlp_engine = LoadedSpacyNlpEngine(loaded_spacy_model = nlp)
 
 
analyzer = AnalyzerEngine(nlp_engine = loaded_nlp_engine)
 
 
def mask_pii(text: str) -> str:
    if not text:
        return text
 
    # Mask Names using Presidio
    #print("Masking PII data...")
    #print(f"Original Text: {text}")
    results = analyzer.analyze(text=text, entities=["PERSON"], language="en")
    #Sort by start index descending to avoid shifting issues
    results = sorted(results, key=lambda x: x.start, reverse=True)
    #print(f"Presidio Results: {results}")
    for r in results:
        start, end = r.start, r.end
        text = text[:start] + "[NAME_REDACTED]" + text[end:]
   
 
    # Step 2: Mask Emails using regex
    email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"
    text = re.sub(email_pattern, "[EMAIL_REDACTED]", text)
 
    # Step 3: Mask Phone Numbers only if + followed by >7 digits
    phone_pattern = r"\+\d{8,}[\d\s\-\(\)]*"
    text = re.sub(phone_pattern, "[PHONE_REDACTED]", text)
    #print(f"{text}")
    return text
 
def mask_text(text, company_patterns=None):
    if not text:
        return text
    text = CREDIT_CARD_PATTERN.sub("[CREDIT_CARD_REDACTED]", text)
    text = BANK_ACC_PATTERN.sub("[BANK_DETAILS_REDACTED]", text)
    text = PHONE_PATTERN.sub("[PHONE_REDACTED]", text)
    text = EMAIL_PATTERN.sub("[EMAIL_REDACTED]", text)
    if company_patterns:
        for cp in company_patterns:
            text = cp.sub("[COMPANY_REDACTED]", text)
   
    # calling this function that masks "Names" in the context
    text_final = mask_pii(text)
 
 
    # with open("Teams-transcript-masking.txt", "w", encoding="utf-8") as file:
    #     file.write(text_final)
    # print(text_final)
    return text_final
 
 
def _apply_mask_to_element(element, company_patterns):
   """
   Applies the mask_text function's logic to an element (Paragraph)
   by clearing its runs and inserting a new run with the masked text.
   This prevents the complete destruction of the parent element's XML,
   which helps preserve non-text elements like images and shapes.
   """
   full_text = element.text
   if not full_text:
       return
   masked_text = mask_text(full_text, company_patterns)
   if full_text != masked_text:
       element._element.clear_content()
       element.add_run(masked_text)
 
 
# This should be fed the file path
def mask_docx(file_path, company_names=None):
   base, ext = os.path.splitext(file_path)
#    output_path = base + "_MASKED_" + ext
   doc = Document(file_path)
   company_patterns = [re.compile(rf"\b{re.escape(name)}\b", re.IGNORECASE) for name in (company_names or [])]
   
   for p in doc.paragraphs:
       _apply_mask_to_element(p, company_patterns)
   for t in doc.tables:
       for row in t.rows:
           for cell in row.cells:
               for p in cell.paragraphs:
                   _apply_mask_to_element(p, company_patterns)
   
   return doc
 
#================ SUMMARIZATION ==================
 
 
def mask_file(file_path, company_names=None):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        return mask_docx(file_path, company_names)
   
def chunk_text(text, max_chars=9000):
    """Split long text into chunks for summarization."""
    chunks = []
    while len(text) > max_chars:
        split_at = text[:max_chars].rfind(".")
        if split_at == -1:
            split_at = max_chars
        chunks.append(text[:split_at + 1].strip())
        text = text[split_at + 1:]
    if text.strip():
        chunks.append(text.strip())
    return chunks
 
def extract_docx_content(docx):
    """Extract text and images from a .docx file."""
    doc = docx
    all_text = ""
    for para in doc.paragraphs:
        all_text += para.text + "\n"
    return all_text.strip()
 
async def summarize_chunks(chunks):
    tasks = []
    for idx, chunk in enumerate(chunks, start=1):
        prompt = f"""Summarize part {idx} of the document:\n\n{chunk}
        Summarize the following  part {idx} of the document:\n\n{chunk}
        By strictly following these rules:
        A. Summary Content Requirements
        Provide summary of the key points discussed such as the symptoms of the incident or failure, discussion about the potential root cause or resolution steps.
        (Don't add ** )in the summary response.
        B. Strict Rules
        - Do not hallucinate or add any information not found in the original text.
        - Do not include any text or commentary beyond the summary.
        - Keep the tone factual and consistent with the source.
    - Let the summary be a free flow summary with out any headings    
    - Don't mention any people name in the summary
        """
        tasks.append(summarize_text_async(prompt, idx))
    return await asyncio.gather(*tasks)
 
async def summarize_text_async(prompt, idx):
    response = await bedrock_invoke_model_async(prompt)
    return response
 
# -------------------- STEP 4: Combine into Final Summary --------------------
 
async def summarize_docx_with_images(docx):
    # print("Extracting DOCX content...")
    text = extract_docx_content(docx)
 
    text_chunks = chunk_text(text)
    # print(f"Split document into {len(text_chunks)} chunks.")
    chunk_summaries= await asyncio.gather(
        summarize_chunks(text_chunks)
    )
 
    combined_input = f"""
        Below are the partial summaries of the document:
        {chunk_summaries}
        Provide summary of the key points discussed such as the symptoms of the incident or failure, discussion about the potential root cause or resolution steps.
        Also follow to give your summary output by following these rules:
        Summary Length Rules
        maximum length of the summary should be 300 words
        minimum length of the summary can be 20 words
 
        Summarization Rules
        - Don't mention any people name in the summary
        - if the document does not contain valid content to summarize or is empty/only has titles, respond with <INVALID_DOCUMENT>
     """
 
    logger.info("Generating final integrated summary...")
 
    final_response = await bedrock_invoke_model_async(combined_input)
    logger.info(final_response)
    return {
        "final_summary": final_response,
        "chunk_summaries": chunk_summaries
    }
 
#================ READING WEBSITES =================
 
def get_files_recursive(drive_id, headers, folder_id=None):
    files = []
    if folder_id:
        url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{folder_id}/children"
    else:
        url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root/children"
 
    res = requests.get(url, headers=headers).json()
    for item in res.get('value', []):
        if item.get('folder'):
            files.extend(get_files_recursive(drive_id, headers, item['id']))
        elif item['name'].lower().endswith('.docx') and "inc" in item['name'].lower():
            created = datetime.fromisoformat(item['createdDateTime'].replace('Z', '+00:00'))
            modified = datetime.fromisoformat(item['lastModifiedDateTime'].replace('Z', '+00:00'))
            #if created >= three_years_ago or modified >= three_years_ago:
            if created >= last_24_hours or modified >= last_24_hours:
                files.append(item)

    return files
 
# ======== Format to inc+file =========
 
def grabInc(fileList):
    outList = []
    for filejson in fileList:
        fname = filejson['name']
        logger.info(f"File name: {fname}")
        incs = re.findall(r'INC\d{7}', fname)
        if incs.__len__() == 0: pass
        else:
            logger.info(f"====== Found incident: {incs[0]}")
            outList.append( (incs[0], filejson) )
    return outList
           
# ======= DOWNLOAD FILE ============
 
def download_file_from_url(item, site_id, drive_id, download_path):
    item_id = item['id']
    download_url = item.get('@microsoft.graph.downloadUrl')
 
    if not download_url:
        logger.info(f"No download URL for item ID: {item_id}")
        return
 
    # Create a unique filename
    unique_filename = f"{site_id}_{drive_id}_{item_id}"
    file_extension = os.path.splitext(item['name'])[1]  # Preserve original extension
    full_filename = unique_filename + file_extension
 
    os.makedirs(download_path, exist_ok=True)
    file_path = os.path.join(download_path, full_filename)
 
    try:
        response = requests.get(download_url, stream=True, verify=False)
        response.raise_for_status()
 
        with open(file_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
 
        logger.info(f"Downloaded: {full_filename}")
        return "./downloads/" + full_filename
    except Exception as e:
        logger.error(f"Download failed for {full_filename}: {str(e)}")
 
# ============= Formatting Row Item ===============
 
def _format_(inc_id , filejson , summaryResult):
 
    #getci
    try:
        ci = _sql_(f"""select configuration_item from p1p2_incidents where inc_id = '{inc_id}' limit 1;""")[0][0]
    except Exception as e:
        logger.info(f"Error fetching CI for incident {inc_id}: {str(e)}")
        ci = None
   
 
    item = {
            "ci" : ci,
            "chunk" : None ,
            "embedding" : _get_embeddings(summaryResult['final_summary']),
            "knowledge_id" : inc_id ,
            "knowledge_type" : "Teams Transcript" ,
            "source" : "Teams",
            "link" : filejson["webUrl"],
            "summary": summaryResult['final_summary'],
    }
    return item
    # ci,chunk,embedding,knowledge_id,knowledge_type,source,link,summary
 
def get_site(HOSTNAME, SITE_PATH, headers):
    url = f"https://graph.microsoft.com/v1.0/sites/{HOSTNAME}:/{SITE_PATH}"
    res = requests.get(url, headers=headers).json()
    return res
 
def get_drives(site_id, headers):
    url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives"
    res = requests.get(url, headers=headers).json()
    return res.get('value', [])
 
# ============== CALL THIS WITH SITE + DRIVE NAME ==============
 
def teams_helper():

    try:
    # Authenticate and get the access token
        auth_url = f'https://login.microsoftonline.com/{TENANT_ID}/oauth2/v2.0/token'
        
        auth_data = {
            'grant_type': 'client_credentials',
            'client_id': CLIENT_ID,
            'client_secret': CLIENT_SECRET,
            'scope': 'https://graph.microsoft.com/.default'
        }
        
        response = requests.post(auth_url, data=auth_data)
        
        access_token = response.json()['access_token']
        #logger.info(f"Access Token retrieved successfully")
        
        # headers
        headers = {
            'Authorization': f'Bearer {access_token}',
            'Content-Type': 'application/json',
            'Prefer': 'HonorNonIndexedQueriesWarningMayFailRandomly'
        }
        
        all_files = []

        parsed = urlparse(site_url[0])
        paths = parsed.path.strip("/").split("/")
        site_path = "/".join(paths[:2])
        host_name = HOST_NAME

        logger.info(f"Processing SharePoint Site path : {site_path}")
        site = get_site(host_name, site_path, headers)
        site_id = site['id']
        logger.info(f"Processing site: {site['name']}")
        drive_id = get_drives(site_id, headers)[0]['id']

        files = get_files_recursive(drive_id, headers)
        logger.info(f"Found {len(files)} files in drive '{drive_id}' of site '{site_id}'")
        all_files.extend(files)
    
        logger.info(f"Total files found (last 24 hrs): {len(all_files)}")
    

        conn = psycopg2.connect(**db_config)
        cur = conn.cursor()

        select_query = """
            SELECt count(knowledge_id) from knowledge_vec where source= 'Teams';
            """
        
        cur.execute(select_query)
        count = cur.fetchall()
        logger.info(f"Pre-check: Total records in knowledge_vec before processing: {count}")



        # for each file
        data = []
        comp = 0
        invalid = 0
        restricted = 0
        for inc_id , filejson in tqdm(grabInc(all_files)):
            # if inc_id in ["INC0338871","INC0333261","INC0372826",'INC0372285','INC0369715','INC0367840','INC0345423','INC0339284','INC0165022','INC0142325','INC0140966','INC0139818','']:
            #     print(f"Processing incident: {inc_id} , file: {filejson}")
            logger.info(f"{comp} Completed , {invalid} Invalid docs , {restricted} Restricted docs ")
            # logger.info(inc_id ," - ", filejson['name'], " | URL:  ",filejson['webUrl'])
            try:
                check_query = f"""
                            SELECT EXISTS (
                                SELECT 1 FROM knowledge_vec WHERE knowledge_id = '{inc_id}')
                            """
                cur.execute(check_query)
                result = cur.fetchone()[0]
                if result:
                    logger.info(f"Record with knowledge_id {inc_id} already exists. Skipping the file.")
                    comp += 1
                    continue


                filepath = download_file_from_url(filejson, site_id, drive_id , "./downloads")
                # Call the summary
                #summary = {"final_summary": " Final summary here ","chunk_summaries": " chunk summaries here "}
                #========== UNCOMMENT ===============
                
                try:
                    doc = mask_file(filepath , ["Easy Jet"]) # UNCOMMENT AFTER DOCX IS THERE
                except Exception as e:
                    logger.error(f"Error masking file {filejson} : {e}")
                    restricted += 1
                    continue

                summary = asyncio.run(summarize_docx_with_images(doc))
                #====================================
                if "<INVALID_DOCUMENT>" in summary['final_summary']:
                    logger.info(f"Skipping invalid document for incident {inc_id} , {filejson['name']} , URL: ,{filejson['webUrl']}")
                    invalid += 1
                    continue
                logger.info(f"Summary for incident {inc_id} :\n {summary['final_summary']}\n\n")
                # format into entry item
                pushRow = _format_(inc_id , filejson , summary)

                #print(json.dumps(pushRow,indent=3))

                # push to table
                conn = psycopg2.connect(**db_config)
                cur = conn.cursor()
            
                insert_query = """
                INSERT INTO knowledge_vec (
                    ci,
                    chunk,
                    embedding,
                    knowledge_id,
                    knowledge_type,
                    source,
                    link,
                    summary
                )
                VALUES (
                    %(ci)s,
                    %(chunk)s,
                    %(embedding)s,
                    %(knowledge_id)s,
                    %(knowledge_type)s,
                    %(source)s,
                    %(link)s,
                    %(summary)s
                )
                """

                cur.execute(insert_query, pushRow)
                conn.commit()
                comp += 1
        
            except Exception as e:
                logger.error(f"Error: {traceback.format_exc()}")
                logger.error(filejson)
                raise
        
        if conn:
            cur.close()
            conn.close()
        # return all_files
    except Exception as e:
        logger.error(f"Error in teams_helper: {traceback.format_exc()}")
        raise


# teams_helper("https://easyjet.sharepoint.com/teams/ServiceIntegration/Shared%20Documents/Forms/AllItems.aspx?id=%2Fteams%2FServiceIntegration%2FShared%20Documents%2FCall%20Recordings&viewid=ccd99ed0%2D2b35%2D411b%2D8f5d%2D988d07287b2b")